# -*- coding: utf-8 -*-
'''
@FileName	:   fout.py
@Created    :   2021/08/07
@Updated    :   2022/06/25
@Author		:   Teddy, goonhope@gmail.com, Zhuhai
@Function	:   word标记，PDF图片批量插入word，PDF转图片等
@notes      ：  知乎公开函数库
'''

import os,time,shutil,zipfile
from functools import wraps
from PIL import Image
from docx import Document
from docx.shared import RGBColor,Inches,Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as align,WD_BREAK
from PIL import Image, ImageDraw, ImageFont
from win32com.client import Dispatch
from Project.function import show,get_excel



def shower(show=True):
    '''print装饰器'''
    def inner(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            print(f"@{func.__name__}...") if show else None
            goal = func(*args, **kwargs)
            print(goal) if len(goal[-1]) < 20 else None
            return goal
        return wrapper
    return inner


def string_check(x,enz="",ink=".",out="$$$",stz="",win=True):
    """字符串check  :return: bool逻辑值  :param x: string 字符串;
     :param enz: endswith字符组;  param ink: in keyword 含有字符；  :param out: not in x不含字符;
     :param stz: endswith字符;   :param win: 忽略大小写，全部转小写;"""
    if win: enz,ink,out,stz,x = (z.lower() for z in (enz,ink,out,stz,x))
    enz, stz = (not z or any(x.endswith(y) for y in z.strip().split()) for z in (enz,stz))  # any
    ink, out = (not z or all(y in x if z == ink else y not in x for y in z.strip().split()) for z in (ink,out))
    return enz and stz and ink and out


# @shower()
def get_root_sub(path,enz="img",file=True,**kargs):
    """获取path路径下的 文件名称或文件夹名list:
    :return root-str, files or dirs-[str]"""
    img = 'jpg jpeg tiff tif png bmp'
    enz = img if enz == "img" else enz
    files = [os.path.split(path)[-1]] if os.path.isfile(path) else [x for x in os.listdir(path) if string_check(x,enz,**kargs)]
    root = path if os.path.isdir(path) else os.path.split(path)[0]  # os.path.dirname(path)
    sub = files if file else [x for x in os.listdir(root) if string_check(x, enz, **kargs) and os.path.isdir(os.path.join(root,x))]
    return root, sub


def extracts_imgs(docx,dlt=True):
    """提取docx/pptx/xlsx中图片，docx"""
    enz = "docx pptx xlsx".split()
    inword = "word ppt xl".split()
    info = {x:y for x,y in zip(enz,inword)} # dict(zip(enz,inword))
    root, docxs = get_root_sub(docx," ".join(enz),out="~$")
    for sdocx in docxs:
        docx = os.path.join(root,sdocx)
        docxb = "_Bakup".join(os.path.splitext(docx))
        shutil.copy(docx,docxb) # bakeup
        nroot = os.path.splitext(docx)[0]
        fzip = nroot + "_.zip"
        os.remove(fzip) if os.path.exists(fzip) else None # exist delete
        os.renames(docxb,fzip)
        xroot = nroot + "_old"
        os.makedirs(xroot) if not os.path.exists(xroot) else None # not exist mk
        with zipfile.ZipFile(fzip, 'r') as fi:
            for file in fi.namelist():
                fi.extract(file, xroot)  # 将压缩包里的word文件夹解压出来
        for x in enz:
            if sdocx.endswith(x):
                media = f"{info[x]}\media"
                imgf = os.path.join(xroot, media)
                shutil.rmtree(nroot) if os.path.exists(nroot) and dlt else None
                os.renames(imgf, nroot) if os.path.exists(imgf) else None # 拷贝到新目录，名称为word文件的名字
                re_name(nroot)
                emf_png(nroot)
        shutil.rmtree(xroot)
        os.remove(fzip)
    os.system(f"""start "" {root}""")


def emf_png(path,dlt=True):
    '''emf图片转png,dlt为真删除emf图片'''
    root,files = get_root_sub(path,enz=".emf")
    for file in files:
        file = os.path.join(root,file)
        Image.open(file).save(file.replace(".emf", ".png"))
        os.remove(file) if dlt else None
    root, files = get_root_sub(path)


def re_name(new_imagedir,st=1):
    '''Windows序号命名修改'''
    # 重命名文件
    l = os.listdir(new_imagedir)
    l_sorted = sorted(l, key=lambda x: int(str(x).split(".")[0].split("e")[-1]))  # 排序转换
    j_s = len(l_sorted)
    for j,pic in enumerate(l_sorted):  # 重命名
        Olddir = os.path.join(new_imagedir, pic)
        if os.path.isfile(Olddir):
            Newdir = os.path.join(new_imagedir, str(j + st).zfill(3) + "_" + pic)
            os.rename(Olddir, Newdir)
    return


def pdf_to_images(path, dpi=300,ro=0,sep="_",ink=".",fz=False,**kwargs):
    '''pdf转为图片，默认DPI=300  需要wand库支持
    path:支持文件名及路径'''
    if fz: pdf_imgs(path,dpi,ro,ink=ink,sep=sep,**kwargs)
    from wand.image import Image
    from wand.color import Color
    path, pdf_names = get_root_sub(path, enz="PDF pdf",ink=ink,**kwargs)
    for pdf_name in pdf_names:
        fpdf = os.path.join(path, pdf_name)
        image_jpeg = Image(filename=fpdf, resolution=dpi)
        pages = len(image_jpeg.sequence)
        for i,img in enumerate(image_jpeg.sequence):
            with Image(image=img) as img_page:
                img_page.format = 'png'
                img_page.background_color = Color('white')
                img_page.alpha_channel = 'remove'
                if ro: img_page.rotate(ro)  # 默认逆时针翻转
                fimg = os.path.splitext(fpdf)[0] + f"{sep + str(i).zfill(3)}.jpg"
                with open(fimg, 'wb') as img_name:
                    img_name.write(img_page.make_blob('jpg'))
        print("@Filename:\t{0}-->Imags:\t{1}".format(pdf_name, pages))


def pdf_imgs(path,dpi=300,ro=0,ink=".",sep="_",**kwargs):
    """可转多图层pdf or svg pdf
    https://pymupdf.readthedocs.io/en/latest/faq.html#faq
    """
    import fitz
    path, pdf_names = get_root_sub(path, enz="PDF pdf",ink=ink,**kwargs)
    for pdf_name in pdf_names:
        fpdf = os.path.join(path,pdf_name)
        pdf = fitz.open(fpdf)
        for page in pdf:
            rotation = page.rotation
            if rotation != 0: page.set_rotation(rotation)
            if ro: page.set_rotation(ro)
            pix = page.get_pixmap(alpha=False,dpi=dpi) #matrix=mat,mat = fitz.Matrix(s, s)
            oimg = os.path.splitext(fpdf)[0] + f"{sep}{str(page.number).zfill(3)}.jpg"
            pix.pil_save(oimg, optimize=True)
            # pix.save(oimg)
        print(f"@{pdf_name}->Img:{pdf.page_count}张，Ppf元信息：", *pdf.metadata.values())
        pdf.close()


def dos(cstr,show=False):
    """运行dos命令"""
    cstr = f'chcp 65001 && {cstr}'  # cmd字符代码UTF8 ，window 默认936-GBK
    result = os.popen(cstr).read().strip()
    print(result) if show else None
    return


def log(goal,fn="_info.log"):
    ''''记录'''
    file = os.path.splitext(__file__)[0] + fn
    with open(file, "a", encoding="utf-8") as f:
        info = "\n".join(["\t".join(x) for x in goal])
        print(info, file=f)
    return True if os.path.exists(file) else False


class Word(object):
    '''word-docx文件 生成类'''
    def __init__(self,template=r"D:\temp\2022\template.docx",clear=True):
        self.template = template
        self.clear = clear
        self.out = "_插图".join(os.path.splitext(self.template))
        self.doc = Document(template)
        self.info = []
        if self.clear: self.doc._body.clear_content()  # 清楚内容

    def add_para(self,content="",style="表图"):
        ''''添加段落：默认空行'''
        return self.doc.add_paragraph(content, style=style)

    def insert_imgs(self,root,imgs,inch=6,sp=True):
        ''''批量插入图片'''
        run = self.add_para().add_run()  # text=None, style=None
        for img in imgs:
            img = os.path.join(root,img)
            width, height = Image.open(img).size
            finch = Inches(inch) if width >= 21 / 2.54 * 72 else None
            run.add_picture(img, width=finch)
        run.add_break(WD_BREAK.PAGE) if sp else None  # 加分页符
        return run

    def del_imgs(self,dir,**kwargs):
        '''批量删除图片文件'''
        root,imgs = get_root_sub(dir,**kwargs)
        for img in imgs:
            os.remove(os.path.join(root,img))

    def replace(self,old="认定",new="技改"):
        '''替换word中字符''' # doc.sections[0].header.paragraphs[0].text
        for para in self.doc.paragraphs:
            if old in para.text:
                para.text = para.text.replace(old, new)
        for sect in self.doc.sections:  # 节
            sect.different_first_page_header_footer = 0 # 取消首页不同
            hf = sect.even_page_header.paragraphs + sect.header.paragraphs
            for hp in hf:
                if old in hp.text:
                    print(hp.text,end="\t")
                    hp.text = hp.text.replace(old, new)
                    print(hp.text)
        self._save()

    def _save(self):
        '''保存文件，如有原重名doc\docx文件先删除'''
        os.remove(self.out) if os.path.exists(self.out) else None
        out2 = self.out.strip("x")
        os.remove(out2) if os.path.exists(out2) else None
        self.doc.save(self.out)
        show(os.path.split(self.template)[0])

    def pdf_word(self,imgdir,pdf=False,ro=0,level=3,dlt=True):
        ''''以目录内pdf文件名为标题，批量插入（pdf转）图片到word文件中'''
        titles = [x.lower().replace(".pdf","") for x in get_root_sub(imgdir,enz=".pdf .PDF")[-1]]
        if pdf: pdf_to_images(imgdir,ro=ro,fz=True)
        for name in titles:
            self.add_para(name , style=f'Heading {str(level)}')
            root, imgs = get_root_sub(imgdir, ink=name,win=True)
            print(name,len(imgs))
            self.insert_imgs(root,imgs,sp=True if name != list(titles)[-1] else False)  # 最后不分页
            fname = os.path.join(root,imgs[0][:imgs[0].rfind("_")] + ".pdf")
            if os.path.exists(fname) and dlt: self.del_imgs(imgdir, ink=name)

    def pdf_order(self,imgdir,kys,pdf=False,ro=0,level=3,dlt=False):
        ''''插入制定文字的图片或pdf文件'''
        if pdf: pdf_to_images(imgdir,ro=ro)
        kys = kys.strip().split()
        for name in kys:
            self.add_para(name , style=f'Heading {str(level)}')
            root, imgs = get_root_sub(imgdir, ink=name)
            print(name,len(imgs),*imgs)
            self.insert_imgs(root,imgs,sp=True if name != kys[-1] else False)  # 最后不分页
            # fname = os.path.join(root,imgs[0][:imgs[0].rfind("_")] + ".pdf")
            # if os.path.exists(fname) and dlt: self.del_imgs(imgdir,ink=name)

    def sub(self,imgdir,pdf=False,level=2,reverse=False,dt=False,dlt=False):
        '''以子文件夹为标题，批量插入子文件中（pdf转）图片到word文件中'''
        dirs = [x for x in os.listdir(imgdir) if os.path.isdir(os.path.join(imgdir,x))]
        if reverse: dirs = sorted(dirs,key=lambda x: int(x[:x.find("、")]))
        for n,dir in enumerate(dirs):
            fdir = os.path.join(imgdir,dir)
            title = dir.strip("0123456789-") if dt else dir
            self.add_para(title, style=f'Heading {str(level)}')
            if pdf: pdf_to_images(fdir)
            root, imgs = get_root_sub(fdir)
            print(dir, len(imgs), imgs)
            self.info.append([str(n).zfill(3), dir, str(len(imgs)),"\t".join(imgs)])
            if imgs:
                self.insert_imgs(root, imgs, sp=True if dir != list(dirs)[-1] else False)  # 最后不分页
                if dlt: self.del_imgs(os.path.join(imgdir,dir))
            else:
                n = level + 1
                fsdir = os.path.join(fdir)
                if os.path.isdir(fsdir):
                    self.sub(fsdir,level=n,dt=False)

    def direct(self,imgdir,ink=""):
        '''直接导入_1张1标题'''
        kys = ink.strip().split() or get_root_sub(imgdir,ink=ink)[-1]
        for n,name in enumerate(kys):
            self.add_para(os.path.splitext(name)[0], style='Heading 3')
            root, imgs = get_root_sub(imgdir, ink=name.lower())
            print(name,len(imgs), imgs)
            self.insert_imgs(root, imgs)
        self._save()

    def add_all(self,imgdir,name="示例"):
        '''全部导入'''
        root, imgs = get_root_sub(imgdir, ink=".")
        print(imgdir,len(imgs), imgs)
        self.add_para(name, style='Heading 3')
        self.insert_imgs(root, imgs)
        self._save()


def get_mark(file,color=(0,0,0)):
    """获取正文word标记，默认非黑色为标记"""
    openf = Document(file)
    hold = dict()
    for para in openf.paragraphs:
        for run in para.runs:
            c = run.font.color.rgb
            if c and c != RGBColor(*color):
                hold.update({run.text:run.font.color.rgb})
    return hold


def get_run(run,info):
    """获取或赋值run style信息"""
    if not info:
        return run.font.size,run.bold,run.font.color.rgb,run.font.name
    else:
        run.font.size, run.bold, run.font.color.rgb, run.font.name=info
        s = run._element
        s.rPr.rFonts.set(qn("w:eastAsia"),info[-1])


def marked(file,keys):
    """标记word正文特定词语"""
    openf = Document(file)
    for text, c in keys.items():
        for para in openf.paragraphs:
            pt = para.text
            if c and text in pt:
                for run in para.runs:
                    info = get_run(run,0)
                    if text == run.text:
                        run.font.color.rgb = c
                para.text = ""
                pt = pt.replace(text,"$%$" + text + "$%$").split("$%$")
                for tex in pt:
                    run = para.add_run(tex)
                    get_run(run,info)
                    if tex == text:
                        run.font.color.rgb = c
    fout = "_ed".join(os.path.splitext(file))
    if os.path.exists(fout): os.remove(fout)
    openf.save(fout)


def pword(dir,pdf=True,dlt=False):
    """主函数"""
    doc = Word()
    doc.pdf_word(dir,pdf,dlt=dlt,ro=0)
    # order = ""
    # doc.pdf_order(dir,order,pdf)
    # doc.sub(dir,True,dt=True))
    # doc.direct(dir)
    # doc.add_all(dir)
    doc._save()

    
def mk_dirs(path,newdirs):
    """批量新建文件夹"""
    for x in newdirs:  # 批量新建文件夹
        x = os.path.join(path,x.strip())
        if not os.path.exists(x): os.makedirs(x)
    return


def get_exts(path,sub=True,dot=False):
    """获取目录内所有文件后缀，默认小写无点"""
    from os.path import splitext
    fn = lambda i: splitext(i)[-1].lower() if dot else splitext(i)[-1].lower().strip(".")
    exts = set(fn(i) for r, d, x in os.walk(path) for i in x if fn(i) and sub or r == path)
    return exts


def filed_by(path,newdirs="",inside=False,sub=True):
    """默认按文件后缀归类文件"""
    fn = lambda x,y: x.lower() in y.lower() if inside else y.lower().endswith(x.lower())
    dir_names = newdirs.strip().split() or get_exts(path)
    mk_dirs(path,dir_names)
    for en,(root, dir, files) in enumerate(os.walk(path,topdown=False)):
        if sub or root == path:
            for file in files:   # 批量文件归到各类文件夹
                for iname in dir_names:
                    if os.path.split(root)[-1].lower() not in dir_names: # 排除新建目录
                        if fn(iname,file):
                            nf = os.path.join(path, iname, file)
                            nf = nf if not os.path.exists(nf) else f"_{os.path.split(root)[-1]}_{str(en)}".join(os.path.splitext(nf)) # 避免重名
                            file = os.path.join(root,file)
                            os.renames(file,nf)
    return


if __name__ == "__main__":
    # main()
    dir = input("请输入目录：").strip()
    pword(dir,dlt=True)
    # extracts_imgs(dir)
    # pdf_to_images(dir)

